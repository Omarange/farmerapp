"""Ingest Bangla agricultural PDFs/DOCX files into a simple retrieval index."""

import argparse
import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import fitz  # pymupdf
from dotenv import load_dotenv

import google.generativeai as genai
from docx import Document

from rag_utils import (
    embed_texts,
    ensure_dir,
    normalize_bangla_text,
    paragraph_chunks,
    save_index,
)


CROP_CONFIG: List[Dict[str, object]] = [
    {
        "id": "amon_dhan",
        "crop_bn": "রোপা আমন ধান",
        "aliases": ["amon", "dhan", "আমন", "ধান"],
        "synonyms": [
            "রোপা আমন ধান",
            "আমন ধান",
            "ধান",
            "Amon Dhan",
            "Aman paddy",
            "paddy",
        ],
        "category": "cereal",
    },
    {
        "id": "tomato",
        "crop_bn": "টমেটো",
        "aliases": ["tomato", "টমেটো"],
        "synonyms": ["টমেটো", "টমেটো চাষ", "Tomato", "Lycopersicon", "tomato farming"],
        "category": "vegetable",
    },
    {
        "id": "watermelon",
        "crop_bn": "তরমুজ",
        "aliases": ["watermelon", "তরমুজ"],
        "synonyms": ["তরমুজ", "Watermelon", "তরমুজ চাষ"],
        "category": "fruit",
    },
    {
        "id": "puishak",
        "crop_bn": "পুঁইশাক",
        "aliases": ["pui", "puishak", "পুঁই", "পুঁইশাক"],
        "synonyms": ["পুঁইশাক", "Pui shak", "Basella", "Potherb"],
        "category": "leafy",
    },
    {
        "id": "begun",
        "crop_bn": "বেগুন",
        "aliases": ["begun", "brinjal", "eggplant", "বেগুন"],
        "synonyms": ["বেগুন", "Brinjal", "Eggplant", "বেগুন চাষ"],
        "category": "vegetable",
    },
    {
        "id": "morich",
        "crop_bn": "মরিচ",
        "aliases": ["morich", "chili", "মরিচ"],
        "synonyms": ["মরিচ", "Chili", "Hot pepper", "কাঁচা মরিচ"],
        "category": "vegetable",
    },
    {
        "id": "shosha",
        "crop_bn": "শসা",
        "aliases": ["shosha", "shasha", "cucumber", "শসা"],
        "synonyms": ["শসা", "শশা", "Cucumber", "শসা চাষ"],
        "category": "vegetable",
    },
]


STAGE_KEYWORDS: Dict[str, List[str]] = {
    "seedbed": ["বীজতলা", "নার্সারি", "অঙ্কুরণ", "seedbed", "nursery"],
    "transplant": ["রোপণ", "চারা রোপণ", "transplant", "রোপাই"],
    "vegetative": ["বৃদ্ধি", "পাতা", "vegetative", "পরিচর্যা"],
    "flower": ["ফুল", "শিষ", "flower"],
    "fruit": ["ফল", "ফল গঠন", "fruit"],
    "harvest": ["কাটা", "সংগ্রহ", "harvest"],
}


TOPIC_KEYWORDS: Dict[str, List[str]] = {
    "fertilizer": ["সার", "ডোজ", "ডিএপি", "ইউরিয়া", "fertilizer"],
    "water": ["সেচ", "পানি", "drainage", "সেচব্যবস্থা"],
    "pest": ["রোগ", "পোকা", "কীট", "কীটনাশক", "পোকামাকড়", "disease"],
    "weather": ["আবহাওয়া", "বৃষ্টি", "তাপমাত্রা", "weather", "আর্দ্রতা"],
    "soil": ["মাটি", "pH", "মাটির", "soil"],
    "variety": ["জাত", "উন্নত জাত", "বারি", "হাইব্রিড", "variety", "cultivar", "জাতসমূহ"],
}


def slugify(value: str) -> str:
    out = re.sub(r"[^0-9a-zA-Z\u0980-\u09FF]+", " ", value)
    out = re.sub(r"\s+", "_", out).strip("_")
    return out.lower()


def detect_crop_info(path: Path) -> Dict[str, object]:
    stem = path.stem
    stem_lower = stem.lower()
    for cfg in CROP_CONFIG:
        aliases = cfg.get("aliases", [])
        for alias in aliases:
            alias_lower = alias.lower()
            if alias_lower and alias_lower in stem_lower:
                return cfg
            if alias and alias in stem:
                return cfg
    # fallback generic config
    return {
        "id": slugify(stem) or "unknown",
        "crop_bn": stem,
        "aliases": [],
        "synonyms": [stem, stem_lower],
        "category": "unknown",
    }


def detect_tags(text: str, keyword_map: Dict[str, List[str]]) -> Set[str]:
    tags: Set[str] = set()
    if not text:
        return tags
    for tag, tokens in keyword_map.items():
        for token in tokens:
            if token and token in text:
                tags.add(tag)
                break
    return tags



def _build_record(
    source: Path,
    page_idx: int,
    chunk: str,
    crop_info: Dict[str, object],
    heading: Optional[str] = None,
    section_id: Optional[str] = None,
) -> dict:
    chunk_norm = normalize_bangla_text(chunk)
    context_for_tags = chunk_norm + "\n" + normalize_bangla_text(heading or "")
    stage_tags = sorted(detect_tags(context_for_tags, STAGE_KEYWORDS))
    topic_tags = sorted(detect_tags(context_for_tags, TOPIC_KEYWORDS))

    bn_synonyms: Set[str] = set()
    roman_synonyms: Set[str] = set()
    for syn in crop_info.get("synonyms", []):
        syn_str = str(syn).strip()
        if not syn_str:
            continue
        if re.search(r"[\u0980-\u09FF]", syn_str):
            norm_syn = normalize_bangla_text(syn_str)
            if norm_syn:
                bn_synonyms.add(norm_syn)
                for part in re.split(r"[\s/]+", norm_syn):
                    part = part.strip()
                    if len(part) >= 2:
                        bn_synonyms.add(part)
        else:
            low_syn = syn_str.lower()
            if low_syn:
                roman_synonyms.add(low_syn)
                for part in re.split(r"[\s/]+", low_syn):
                    part = part.strip()
                    if len(part) >= 3:
                        roman_synonyms.add(part)

    if heading:
        head_norm = normalize_bangla_text(heading)
        if head_norm:
            for part in re.split(r"[\s/]+", head_norm):
                part = part.strip()
                if part:
                    if re.search(r"[\u0980-\u09FF]", part):
                        bn_synonyms.add(part)
                    else:
                        roman_synonyms.add(part.lower())

    meta = {
        "crop_id": crop_info.get("id", "unknown"),
        "crop_bn": crop_info.get("crop_bn", source.stem),
        "category": crop_info.get("category", "unknown"),
        "synonyms": sorted(bn_synonyms) + sorted(roman_synonyms),
        "stage_tags": stage_tags,
        "topic_tags": topic_tags,
        "heading": heading,
        "section_id": section_id,
    }

    return {
        "source": source.name,
        "page": page_idx,
        "text": chunk_norm,
        "meta": meta,
    }


def extract_pdf(path: Path, crop_info: Dict[str, object]) -> List[dict]:
    records: List[dict] = []
    with fitz.open(path) as doc:
        for page_idx, page in enumerate(doc, start=1):
            text = page.get_text("text")
            text = normalize_bangla_text(text)
            if not text:
                continue
            for chunk in paragraph_chunks(text):
                records.append(_build_record(path, page_idx, chunk, crop_info))
    return records


def extract_docx(path: Path, crop_info: Dict[str, object]) -> List[dict]:
    records: List[dict] = []
    doc = Document(path)

    sections: List[Tuple[str, List[str]]] = []
    current_heading: Optional[str] = None
    buffer: List[str] = []

    def flush_section():
        nonlocal sections, buffer, current_heading
        if buffer:
            sections.append((current_heading, list(buffer)))
        buffer = []

    for para in doc.paragraphs:
        raw = para.text.strip()
        if not raw:
            continue
        text = normalize_bangla_text(raw)
        if not text:
            continue

        style = (getattr(para.style, "name", "") or "").lower()
        is_heading = False
        if style.startswith("heading") or style in {"title", "subtitle"}:
            is_heading = True
        elif len(raw) < 70 and raw.endswith(":"):
            is_heading = True
        elif len(raw.split()) <= 6 and raw.isupper():
            is_heading = True

        if is_heading:
            flush_section()
            current_heading = text.rstrip(":")
            continue

        buffer.append(text)

    flush_section()

    if not sections:
        text = "\n".join(normalize_bangla_text(p.text) for p in doc.paragraphs if p.text.strip())
        for idx, chunk in enumerate(paragraph_chunks(text), start=1):
            records.append(_build_record(path, idx, chunk, crop_info))
        return records

    for sec_idx, (heading, paras) in enumerate(sections, start=1):
        combined = "\n".join(paras)
        chunks = paragraph_chunks(combined, max_chars=750, min_chars=160)
        if not chunks:
            continue
        for chunk_idx, chunk in enumerate(chunks, start=1):
            records.append(
                _build_record(
                    path,
                    page_idx=sec_idx,
                    chunk=chunk,
                    crop_info=crop_info,
                    heading=heading,
                    section_id=f"{sec_idx}-{chunk_idx}",
                )
            )
    return records


def ingest(data_dir: Path, out_dir: Path) -> None:
    input_paths = sorted(
        [p for p in data_dir.glob("**/*.pdf") if p.is_file()] +
        [p for p in data_dir.glob("**/*.docx") if p.is_file()]
    )
    if not input_paths:
        raise SystemExit(f"No PDF or DOCX files found under {data_dir}")

    all_records: List[dict] = []
    for path in input_paths:
        print(f"Extracting {path}…")
        crop_info = detect_crop_info(path)
        if path.suffix.lower() == ".pdf":
            all_records.extend(extract_pdf(path, crop_info))
        else:
            all_records.extend(extract_docx(path, crop_info))

    if not all_records:
        raise SystemExit("No text extracted from documents.")

    print(f"Embedding {len(all_records)} chunks…")
    texts = [rec["text"] for rec in all_records]
    embeddings = embed_texts(texts)

    ensure_dir(out_dir)
    save_index(embeddings, all_records, out_dir=out_dir)
    print(f"Saved index to {out_dir}")


def main():
    parser = argparse.ArgumentParser(description="Build Bangla RAG index from PDFs")
    parser.add_argument("--data-dir", default="data/pdfs", help="Directory containing PDF/DOCX files")
    parser.add_argument(
        "--out-dir",
        default="data",
        help="Directory to write embeddings/metadata (defaults to data/)",
    )
    args = parser.parse_args()

    load_dotenv()

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise SystemExit("GEMINI_API_KEY must be set in environment or .env")
    genai.configure(api_key=api_key)

    data_dir = Path(args.data_dir)
    out_dir = Path(args.out_dir)

    ingest(data_dir, out_dir)


if __name__ == "__main__":
    main()
